"""
Security Assessment Report (SAR) generator -- OSCAL assessment-results + Word.

Summarizes the assessment of a system's controls: which are satisfied vs other-
than-satisfied, with findings drawn from (a) control responses not fully
implemented / unreviewed and (b) open STIG findings. Completes the core RMF set
alongside the SSP, SAP (test plans), and POA&M.

Same OSCAL caveat as ssp/poam: emits the core structure; validate with
oscal-pydantic before relying on platform import.

OSCAL version targeted: 1.1.2
"""

from __future__ import annotations

import datetime as _dt
import json
import uuid
from pathlib import Path

OSCAL_VERSION = "1.1.2"
# A reviewed control is satisfied whether implemented directly or inherited from a CCP.
_SATISFIED = {"implemented", "inherited"}


def _now() -> str:
    return _dt.datetime.now(_dt.timezone.utc).isoformat()


def assess(conn, system_id: str, catalog_version_id: str) -> dict:
    """Return assessment summary + findings for a system."""
    sysrow = conn.execute("SELECT system_id, name FROM systems WHERE system_id=?",
                          (system_id,)).fetchone()
    if sysrow is None:
        raise ValueError(f"no such system: {system_id}")
    reqs = conn.execute(
        """SELECT control_id, status, needs_review, statement, origin
             FROM implemented_requirements
            WHERE system_id=? AND catalog_version_id=? ORDER BY control_id""",
        (system_id, catalog_version_id)).fetchall()
    inherited = sum(1 for r in reqs if r["origin"] == "inherited")

    findings = []
    satisfied = 0
    for r in reqs:
        ok = (r["status"] in _SATISFIED) and not r["needs_review"]
        if ok:
            satisfied += 1
        else:
            reason = "awaiting review" if r["needs_review"] else f"status: {r['status'] or 'planned'}"
            findings.append({
                "control_id": r["control_id"], "result": "other-than-satisfied",
                "source": "control-response",
                "description": f"{r['control_id'].upper()} not satisfied ({reason})."})

    # STIG open findings
    try:
        from . import stig as _stig
        for f in _stig.open_findings(conn, system_id):
            findings.append({"control_id": f.get("control_id", ""),
                             "result": "other-than-satisfied", "source": "stig",
                             "description": f["weakness"]})
    except Exception:
        pass

    # NIST 800-53A assessment objectives + methods per assessed control
    objectives_by_control: dict[str, list[str]] = {}
    total_objectives = 0
    methods: set[str] = set()
    try:
        from . import assessment
        for r in reqs:
            cid = r["control_id"]
            objs = assessment.objectives(conn, cid, catalog_version_id)
            if objs:
                objectives_by_control[cid] = [o.id for o in objs if o.id]
                total_objectives += len(objs)
                for m in assessment.methods(conn, cid, catalog_version_id):
                    methods.add(m.method)
    except Exception:
        pass

    return {"system_id": sysrow["system_id"], "system_name": sysrow["name"],
            "assessed": len(reqs), "satisfied": satisfied,
            "other_than_satisfied": len(reqs) - satisfied, "findings": findings,
            "inherited": inherited,
            "objectives_by_control": objectives_by_control,
            "total_objectives": total_objectives,
            "methods": sorted(methods)}


_METHOD_LABEL = {"EXAMINE": "Examine", "INTERVIEW": "Interview", "TEST": "Test"}


def build_oscal_sar(conn, *, system_id: str, catalog_version_id: str) -> dict:
    a = assess(conn, system_id, catalog_version_id)
    result_uuid = str(uuid.uuid4())
    obj_by_ctrl = a.get("objectives_by_control", {})

    def _target_id(cid: str) -> str:
        ids = obj_by_ctrl.get(cid)
        return ids[0] if ids else f"{(cid or 'finding').lower()}_obj"

    oscal_findings = [{
        "uuid": str(uuid.uuid4()),
        "title": f"{(f['control_id'] or 'finding').upper()} — {f['result']}",
        "description": f["description"],
        "target": {
            "type": "objective-id",
            "target-id": _target_id(f["control_id"]),
            "status": {"state": "not-satisfied", "reason": "fail"},
        },
        "props": [
            {"name": "control-id", "value": (f["control_id"] or "").lower(),
             "ns": "https://complyforge.local/ns/oscal"},
            {"name": "finding-source", "value": f["source"],
             "ns": "https://complyforge.local/ns/oscal"},
        ],
    } for f in a["findings"]]

    return {"assessment-results": {
        "uuid": str(uuid.uuid4()),
        "metadata": {
            "title": f"Security Assessment Report — {a['system_name']}",
            "last-modified": _now(), "version": "0.1-draft",
            "oscal-version": OSCAL_VERSION,
            "props": [{"name": "generated-by", "value": "ComplyForge"}],
        },
        "import-ap": {"href": f"#sap_{system_id}"},
        "results": [{
            "uuid": result_uuid,
            "title": "Control Assessment Results",
            "description": (
                f"{a['assessed']} controls assessed; {a['satisfied']} satisfied "
                f"({a.get('inherited', 0)} inherited from a common control provider), "
                f"{a['other_than_satisfied']} other-than-satisfied. "
                + (f"Assessed against {a['total_objectives']} NIST SP 800-53A "
                   f"assessment objectives (methods: "
                   f"{', '.join(_METHOD_LABEL.get(m, m) for m in a['methods']) or '—'})."
                   if a.get("total_objectives") else "")),
            "start": _now(),
            "reviewed-controls": _reviewed_controls(a),
            "findings": oscal_findings,
        }],
    }}


def _reviewed_controls(a: dict) -> dict:
    """Reviewed controls + the specific 800-53A objectives they were assessed against."""
    rc: dict = {"control-selections": [{"include-all": {}}]}
    all_obj_ids = [oid for ids in a.get("objectives_by_control", {}).values()
                   for oid in ids]
    if all_obj_ids:
        rc["control-objective-selections"] = [{
            "description": "NIST SP 800-53A assessment objectives for the assessed controls.",
            "include-objectives": [{"objective-id": oid} for oid in all_obj_ids],
        }]
    return rc


def write_oscal_sar(sar: dict, path: str | Path) -> Path:
    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(sar, indent=2))
    return path


def write_word_sar(conn, *, system_id: str, catalog_version_id: str, path: str | Path,
                   prepared_by: str = "", prepared_for: str = "", brand_color: str = "") -> Path:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from ._docx_util import color_run, prepared_block

    a = assess(conn, system_id, catalog_version_id)
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"Security Assessment Report\n{a['system_name']}")
    r.bold = True; r.font.size = Pt(16); color_run(r, brand_color)
    doc.add_paragraph(f"Version 0.1 (DRAFT)   ·   {_dt.date.today():%d %b %Y}")
    prepared_block(doc, prepared_by, prepared_for)

    doc.add_paragraph("1. Executive Summary", style="Heading 1")
    doc.add_paragraph(f"{a['assessed']} controls assessed. {a['satisfied']} satisfied "
                      f"({a.get('inherited', 0)} inherited from a common control provider); "
                      f"{a['other_than_satisfied']} other-than-satisfied. "
                      f"{len(a['findings'])} finding(s) below feed the POA&M.")

    obj_by_ctrl = a.get("objectives_by_control", {})
    if a.get("total_objectives"):
        doc.add_paragraph("2. Assessment Scope (NIST SP 800-53A)", style="Heading 1")
        methods = ", ".join(_METHOD_LABEL.get(m, m) for m in a["methods"]) or "—"
        doc.add_paragraph(
            f"Controls were assessed against {a['total_objectives']} authoritative "
            f"NIST SP 800-53A assessment objectives using the following methods: {methods}.")
        tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Control"
        tbl.rows[0].cells[1].text = "Assessment objectives"
        for cid in sorted(obj_by_ctrl):
            c = tbl.add_row().cells
            c[0].text = cid.upper(); c[1].text = str(len(obj_by_ctrl[cid]))
        findings_heading = "3. Findings"
    else:
        findings_heading = "2. Findings"

    doc.add_paragraph(findings_heading, style="Heading 1")
    if not a["findings"]:
        doc.add_paragraph("No findings — all assessed controls satisfied.")
    else:
        tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
        for i, h in enumerate(("#", "Control", "Source", "Finding")):
            tbl.rows[0].cells[i].text = h
        for n, f in enumerate(a["findings"], 1):
            c = tbl.add_row().cells
            c[0].text = str(n); c[1].text = (f["control_id"] or "").upper()
            c[2].text = f["source"]; c[3].text = f["description"]

    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
