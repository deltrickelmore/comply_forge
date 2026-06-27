"""
Plan of Action & Milestones (POA&M) generator -- OSCAL JSON + Word.

A POA&M tracks open weaknesses and their remediation. ComplyForge derives POA&M
items from a system's control responses that are NOT fully implemented
(status in planned/partial/na, or still needs_review), or from an explicit list of
findings. As with the SSP, the OSCAL output is the CORE structure (not yet
schema-validated -- validate with oscal-pydantic/compliance-trestle before relying
on platform import).

OSCAL version targeted: 1.1.2
"""

from __future__ import annotations

import datetime as _dt
import json
import uuid
from pathlib import Path

OSCAL_VERSION = "1.1.2"
_OPEN_STATUSES = {"planned", "partial", "na", None, ""}


def _now() -> str:
    return _dt.datetime.now(_dt.timezone.utc).isoformat()


def open_items(conn, system_id: str, catalog_version_id: str) -> list[dict]:
    """Control responses that represent open weaknesses (not fully implemented)."""
    rows = conn.execute(
        """SELECT control_id, status, statement, needs_review
             FROM implemented_requirements
            WHERE system_id=? AND catalog_version_id=? ORDER BY control_id""",
        (system_id, catalog_version_id)).fetchall()
    items = []
    for r in rows:
        status = (r["status"] or "").lower()
        if status in _OPEN_STATUSES or r["needs_review"]:
            items.append(dict(r))
    return items


def _collect(conn, system_id, catalog_version_id, findings, include_stig):
    if findings is not None:
        return findings
    src = open_items(conn, system_id, catalog_version_id)
    if include_stig:
        try:
            from . import stig as _stig
            src = src + _stig.open_findings(conn, system_id)
        except Exception:
            pass
    return src


def build_oscal_poam(conn, *, system_id: str, catalog_version_id: str,
                     findings: list[dict] | None = None,
                     include_stig: bool = True) -> dict:
    sysrow = conn.execute("SELECT system_id, name FROM systems WHERE system_id=?",
                          (system_id,)).fetchone()
    if sysrow is None:
        raise ValueError(f"no such system: {system_id}")

    src = _collect(conn, system_id, catalog_version_id, findings, include_stig)
    poam_items = []
    for it in src:
        cid = (it.get("control_id") or "").upper()
        status = it.get("status") or "planned"
        weakness = it.get("weakness") or (
            f"Control {cid} is not fully implemented (status: {status})"
            + ("; draft response awaiting human review." if it.get("needs_review") else "."))
        poam_items.append({
            "uuid": str(uuid.uuid4()),
            "title": f"{cid} weakness",
            "description": weakness,
            "props": [
                {"name": "control-id", "value": (it.get("control_id") or "").lower(),
                 "ns": "https://complyforge.local/ns/oscal"},
                {"name": "poam-status", "value": status,
                 "ns": "https://complyforge.local/ns/oscal"},
            ],
            "remarks": (it.get("statement") or "")[:2000],
        })

    return {"plan-of-action-and-milestones": {
        "uuid": str(uuid.uuid4()),
        "metadata": {
            "title": f"Plan of Action and Milestones — {sysrow['name']}",
            "last-modified": _now(),
            "version": "0.1-draft",
            "oscal-version": OSCAL_VERSION,
            "props": [{"name": "generated-by", "value": "ComplyForge"}],
        },
        "system-id": {"id": sysrow["system_id"]},
        "poam-items": poam_items,
    }}


def write_oscal_poam(poam: dict, path: str | Path) -> Path:
    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(poam, indent=2))
    return path


def write_word_poam(conn, *, system_id: str, catalog_version_id: str,
                    path: str | Path, findings: list[dict] | None = None,
                    include_stig: bool = True, org_name: str = "") -> Path:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sysrow = conn.execute("SELECT name FROM systems WHERE system_id=?",
                          (system_id,)).fetchone()
    src = _collect(conn, system_id, catalog_version_id, findings, include_stig)

    doc = docx.Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"Plan of Action and Milestones (POA&M)\n{sysrow['name'] if sysrow else system_id}")
    r.bold = True; r.font.size = Pt(16)
    doc.add_paragraph(f"Version 0.1 (DRAFT)   ·   {_dt.date.today():%d %b %Y}   ·   "
                      f"{len(src)} open item(s)"
                      + (f"   ·   Prepared for: {org_name}" if org_name else ""))

    cols = ("Item", "Control", "Weakness / Deficiency", "Status",
            "Scheduled Completion", "POC")
    tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Table Grid"
    for i, h in enumerate(cols):
        tbl.rows[0].cells[i].text = h
    for n, it in enumerate(src, 1):
        cid = (it.get("control_id") or "").upper()
        status = it.get("status") or "planned"
        weakness = it.get("weakness") or f"{cid} not fully implemented (status: {status})"
        cells = tbl.add_row().cells
        cells[0].text = str(n); cells[1].text = cid; cells[2].text = weakness
        cells[3].text = status; cells[4].text = ""; cells[5].text = ""

    if not src:
        doc.add_paragraph("No open items — all in-scope controls are fully implemented "
                          "and reviewed.")

    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
