"""
System Security Plan (SSP) generator -- OSCAL JSON + Word.

Assembles a system's implemented_requirements (the control responses) into an SSP.
This is the #1 RMF output. Two renderings:
  * build_oscal_ssp()  -> OSCAL system-security-plan dict (JSON-serializable)
  * write_word_ssp()   -> a human-readable .docx

OSCAL NOTE (honest debt): this emits the CORE OSCAL SSP structure (metadata,
import-profile, system-characteristics, system-implementation, control-implementation
with implemented-requirements/by-components). It is structurally correct but is NOT
yet validated against the OSCAL JSON schema. Before relying on eMASS/Xacta import,
add schema validation via oscal-pydantic or compliance-trestle (validate_oscal()
hooks that in if installed).

AUDIT GUARDRAIL: by default the SSP includes all requirements but flags how many are
unreviewed. Pass require_reviewed=True to refuse export while any draft is unreviewed
(the tool drafts; a human authorizes).

OSCAL version targeted: 1.1.2
"""

from __future__ import annotations

import datetime as _dt
import json
import uuid
from pathlib import Path

OSCAL_VERSION = "1.1.2"
_OBJ = ("confidentiality", "integrity", "availability")


def _now() -> str:
    return _dt.datetime.now(_dt.timezone.utc).isoformat()


def _system(conn, system_id: str) -> dict:
    row = conn.execute(
        "SELECT system_id, name, description, impact_level FROM systems WHERE system_id=?",
        (system_id,)).fetchone()
    if row is None:
        raise ValueError(f"no such system: {system_id}")
    return dict(row)


def _requirements(conn, system_id: str, catalog_version_id: str) -> list[dict]:
    rows = conn.execute(
        """SELECT ir.control_id, ir.status, ir.statement, ir.needs_review,
                  ir.reviewed_by, ir.origin, ps.name AS inherited_from
             FROM implemented_requirements ir
             LEFT JOIN implemented_requirements src ON src.ir_id=ir.source_ir_id
             LEFT JOIN systems ps ON ps.system_id=src.system_id
            WHERE ir.system_id=? AND ir.catalog_version_id=? ORDER BY ir.control_id""",
        (system_id, catalog_version_id)).fetchall()
    return [dict(r) for r in rows]


def build_oscal_ssp(conn, *, system_id: str, catalog_version_id: str,
                    baseline_impact: str = "moderate",
                    require_reviewed: bool = False) -> dict:
    sysrow = _system(conn, system_id)
    reqs = _requirements(conn, system_id, catalog_version_id)
    unreviewed = sum(1 for r in reqs if r["needs_review"])
    if require_reviewed and unreviewed:
        raise ValueError(f"refusing SSP export: {unreviewed} requirement(s) unreviewed. "
                         "A human must approve them (Review Queue) first.")

    impact = (sysrow["impact_level"] or baseline_impact).lower()
    this_system_uuid = str(uuid.uuid4())

    from . import evidence as _ev
    implemented = []
    for r in reqs:
        by_component = {
            "component-uuid": this_system_uuid,
            "uuid": str(uuid.uuid4()),
            "description": r["statement"] or "",
        }
        ev_links = _ev.oscal_links(conn, system_id, r["control_id"])
        if ev_links:
            by_component["links"] = ev_links
        implemented.append({
            "uuid": str(uuid.uuid4()),
            "control-id": r["control_id"],
            "props": [
                {"name": "implementation-status",
                 "value": r["status"] or "planned",
                 "ns": "https://complyforge.local/ns/oscal"},
                {"name": "review-status",
                 "value": "reviewed" if not r["needs_review"] else "draft-needs-review",
                 "ns": "https://complyforge.local/ns/oscal"},
                {"name": "evidence-count", "value": str(len(ev_links)),
                 "ns": "https://complyforge.local/ns/oscal"},
            ],
            "by-components": [by_component],
        })
        if r.get("origin") == "inherited" and r.get("inherited_from"):
            implemented[-1]["props"].append(
                {"name": "inherited-from", "value": r["inherited_from"],
                 "ns": "https://complyforge.local/ns/oscal"})

    ssp = {"system-security-plan": {
        "uuid": str(uuid.uuid4()),
        "metadata": {
            "title": f"System Security Plan — {sysrow['name']}",
            "last-modified": _now(),
            "version": "0.1-draft",
            "oscal-version": OSCAL_VERSION,
            "props": [{"name": "generated-by", "value": "ComplyForge"}],
        },
        "import-profile": {"href": f"#{catalog_version_id.replace('@','_')}_{impact}_baseline"},
        "system-characteristics": {
            "system-ids": [{"id": sysrow["system_id"]}],
            "system-name": sysrow["name"],
            "description": sysrow["description"] or "",
            "security-sensitivity-level": impact,
            "system-information": {"information-types": [{
                "uuid": str(uuid.uuid4()),
                "title": "System information",
                "description": "Information processed, stored, and transmitted by the system.",
                "confidentiality-impact": {"base": impact},
                "integrity-impact": {"base": impact},
                "availability-impact": {"base": impact},
            }]},
            "security-impact-level": {
                "security-objective-confidentiality": impact,
                "security-objective-integrity": impact,
                "security-objective-availability": impact,
            },
            "status": {"state": "operational"},
            "authorization-boundary": {
                "description": f"The authorization boundary of {sysrow['name']} encompasses "
                               "the system components, services, and interconnections documented "
                               "in this SSP. [Refine with the authorization boundary diagram.]"
            },
        },
        "system-implementation": {
            "users": [{"uuid": str(uuid.uuid4()), "title": "System users",
                       "role-ids": ["user"]}],
            "components": [{
                "uuid": this_system_uuid,
                "type": "this-system",
                "title": sysrow["name"],
                "description": sysrow["description"] or "",
                "status": {"state": "operational"},
            }],
        },
        "control-implementation": {
            "description": f"Control implementation for {sysrow['name']} "
                           f"({len(implemented)} controls; {unreviewed} unreviewed).",
            "implemented-requirements": implemented,
        },
    }}
    return ssp


def write_oscal_ssp(ssp: dict, path: str | Path) -> Path:
    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(ssp, indent=2))
    return path


def validate_oscal(ssp: dict) -> tuple[bool, str]:
    """Validate against the official NIST OSCAL SSP schema. Returns (ok, message)."""
    from . import validation
    ok, msgs = validation.validate(ssp, "ssp")
    if ok:
        return True, "OSCAL SSP schema validation passed (NIST v1.1.2)."
    return False, f"{len(msgs)} OSCAL issue(s): " + "; ".join(msgs[:3])


def write_word_ssp(conn, *, system_id: str, catalog_version_id: str,
                   path: str | Path, baseline_impact: str = "moderate",
                   prepared_by: str = "", prepared_for: str = "", brand_color: str = "") -> Path:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sysrow = _system(conn, system_id)
    reqs = _requirements(conn, system_id, catalog_version_id)
    impact = (sysrow["impact_level"] or baseline_impact).lower()

    doc = docx.Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

    from ._docx_util import color_run, prepared_block
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"System Security Plan\n{sysrow['name']}"); r.bold = True; r.font.size = Pt(18)
    color_run(r, brand_color)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Impact Level: {impact.title()}   ·   Version 0.1 (DRAFT — needs review)\n"
                f"{_dt.date.today():%d %b %Y}")
    prepared_block(doc, prepared_by, prepared_for)
    doc.add_page_break()

    doc.add_paragraph("1. System Characteristics", style="Heading 1")
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    for k, v in [("System Name", sysrow["name"]), ("System ID", sysrow["system_id"]),
                 ("Description", sysrow["description"] or ""),
                 ("Security Impact Level", impact.title()),
                 ("Confidentiality / Integrity / Availability", f"{impact.title()} / {impact.title()} / {impact.title()}")]:
        cells = tbl.add_row().cells; cells[0].text = k; cells[1].text = v

    doc.add_paragraph("2. Control Implementation", style="Heading 1")
    unreviewed = sum(1 for x in reqs if x["needs_review"])
    doc.add_paragraph(f"{len(reqs)} controls documented; {unreviewed} awaiting human review.")
    for req in reqs:
        doc.add_paragraph(f"{req['control_id'].upper()}", style="Heading 2")
        flag = "  [DRAFT — NEEDS REVIEW]" if req["needs_review"] else f"  [reviewed by {req['reviewed_by']}]"
        doc.add_paragraph(f"Status: {req['status'] or 'planned'}{flag}")
        doc.add_paragraph(req["statement"] or "[no statement drafted]")

    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
