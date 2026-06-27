"""
RMF Control-Family Plan generator -- the "control plan per family" button.

Reproduces the structure of the AFSV-AFLIS Security Assessment & Authorization
(CA) Plan sample, generalized to ANY 800-53 control family:

  Title page + Revision History
  1. Introduction
     1.1 Purpose            -> Compliance Matrix (the family's controls)
     1.2 Scope
     1.3 Roles and Responsibilities  (table)
     1.4 Government Personnel         (table)
     1.5 Applicable Guidance and Directives
     1.6 Dissemination
     1.7 Review and Authorization
  2. <FAMILY NAME>          (H1)
     2.x <CONTROL TITLE>    (H2) -> narrative + CCI table (AP Acronym|CCI|CCI Definition|Response)

The per-control narrative is LLM-drafted from the control text (or a deterministic
draft without a key). System-specific values come from a SystemProfile (defaults
mirror the sample so output looks right out of the box; override per engagement).
Every plan is a draft for human review.

Built with python-docx (no Node required).
"""

from __future__ import annotations

import datetime as _dt
from dataclasses import dataclass, field
from pathlib import Path

from .llm_provider import LLMProvider, get_provider

FAMILY_TITLES = {
    "ac": "Access Control", "at": "Awareness and Training",
    "au": "Audit and Accountability",
    "ca": "Assessment, Authorization, and Monitoring",
    "cm": "Configuration Management", "cp": "Contingency Planning",
    "ia": "Identification and Authentication", "ir": "Incident Response",
    "ma": "Maintenance", "mp": "Media Protection",
    "pe": "Physical and Environmental Protection", "pl": "Planning",
    "pm": "Program Management", "ps": "Personnel Security",
    "pt": "PII Processing and Transparency", "ra": "Risk Assessment",
    "sa": "System and Services Acquisition",
    "sc": "System and Communications Protection",
    "si": "System and Information Integrity",
    "sr": "Supply Chain Risk Management",
}

NARRATIVE_SYSTEM_PROMPT = """\
You are writing the implementation narrative for one NIST 800-53 control inside a
Security Assessment & Authorization control-family plan for a specific system.
Write 1-3 plain paragraphs describing how the system implements the control and
its enhancements. Be specific to the system described; do not invent capabilities.
Do not attest to compliance, write headings, or restate the control ID. Prose only.
"""


@dataclass
class Personnel:
    role: str
    name: str = ""
    org: str = ""


@dataclass
class SystemProfile:
    system_name: str = ""           # the entity / information system being documented
    system_long_name: str = ""
    enclave: str = ""               # enclave / authorization boundary (optional)
    agency: str = ""                # agency or company the document is being completed for
    baseline: str = "Moderate"
    catalog_version_id: str = "nist_800_53@rev5"
    guidance: list[str] = field(default_factory=lambda: [
        "NIST SP 800-53 Security and Privacy Controls for Information Systems and Organizations",
        "NIST SP 800-37 Risk Management Framework",
        "NIST SP 800-53A Assessing Security and Privacy Controls",
        "FIPS 199 / FIPS 200 Security Categorization and Minimum Requirements",
    ])
    roles: list[tuple[str, str, str]] = field(default_factory=lambda: [
        ("Authorizing Official (AO)", "Government", "Accepts risk and grants authorization to operate."),
        ("Information System Owner (ISO)", "Government", "Responsible for the overall procurement, development, and operation of the system."),
        ("Information System Security Manager (ISSM)", "Government", "Manages the system's cybersecurity program and RMF activities."),
        ("System Administrator", "Privileged", "Maintains the system and implements technical controls."),
        ("User", "General", "Operates the system in accordance with the rules of behavior."),
    ])
    personnel: list[Personnel] = field(default_factory=lambda: [
        Personnel("Information System Owner (ISO)", "", ""),
        Personnel("Information System Security Manager (ISSM)", "", ""),
        Personnel("Information System Security Officer (ISSO)", "", ""),
    ])

    def display_name(self) -> str:
        return self.system_long_name or self.system_name or "[Entity / System Name]"

    def short_name(self) -> str:
        return self.system_name or self.system_long_name or "[Entity / System Name]"

    def description(self) -> str:
        enc = f" operating within the {self.enclave}" if self.enclave else ""
        return f"{self.display_name()}{enc} at the {self.baseline} impact baseline."


def _family_controls(conn, family: str, catalog_version_id: str,
                     baseline_id: str | None) -> list[dict]:
    """Controls in a family (optionally limited to a baseline), ordered by id."""
    fam = family.upper()
    if baseline_id:
        rows = conn.execute(
            """SELECT c.control_id, c.title, c.statement, c.guidance
                 FROM controls c JOIN baseline_controls b
                   ON b.control_id=c.control_id
                WHERE c.catalog_version_id=? AND c.family=? AND b.baseline_id=?
                ORDER BY c.control_id""",
            (catalog_version_id, fam, baseline_id)).fetchall()
        if rows:
            return [dict(r) for r in rows]
    rows = conn.execute(
        """SELECT control_id, title, statement, guidance FROM controls
            WHERE catalog_version_id=? AND family=? ORDER BY control_id""",
        (catalog_version_id, fam)).fetchall()
    return [dict(r) for r in rows]


def _draft_narrative(provider: LLMProvider, ctrl: dict, profile: SystemProfile) -> tuple[str, bool]:
    prompt = (f"Control {ctrl['control_id'].upper()} -- {ctrl['title']}\n"
              f"Requirement: {ctrl['statement']}\n"
              f"Guidance: {ctrl['guidance']}\n\n"
              f"System: {profile.description()}\n\nWrite the implementation narrative.")
    res = provider.complete(system=NARRATIVE_SYSTEM_PROMPT, prompt=prompt, max_tokens=1200)
    text = res.text.strip()
    structured = res.provider != "fake" and bool(text)
    if not structured:
        # deterministic fallback: ground in the control's own text
        text = (f"{ctrl['statement']}\n\n"
                f"{profile.short_name()} implements this control. [Draft narrative — review "
                f"and tailor to the system's actual implementation. {ctrl['guidance'][:300]}]")
    return text, structured


# --------------------------------------------------------------------------- #
# Word writer
# --------------------------------------------------------------------------- #
def _clear_body(doc) -> None:
    """Remove existing body content but keep the final sectPr (headers/footers/margins)."""
    from docx.oxml.ns import qn
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sectPr:
            body.remove(child)


def _table_style(doc) -> str | None:
    names = {s.name for s in doc.styles}
    for cand in ("Table Grid", "Light Grid", "Table Grid Light", "Grid Table 1 Light"):
        if cand in names:
            return cand
    return None


def generate_family_plan(conn, *, family: str, profile: SystemProfile | None = None,
                         baseline_id: str | None = None,
                         provider: LLMProvider | None = None,
                         out_path: str | Path | None = None,
                         template_path: str | Path | None = None,
                         prepared_by: str = "", brand_color: str = "") -> Path:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from ._docx_util import color_run, prepared_block

    provider = provider or get_provider()
    profile = profile or SystemProfile()
    family = family.lower()
    fam_title = FAMILY_TITLES.get(family, family.upper())
    baseline_id = baseline_id or (f"nist_800_53b@{profile.baseline.lower()}"
                                  if profile.baseline else None)

    controls = _family_controls(conn, family, profile.catalog_version_id, baseline_id)
    if not controls:
        raise ValueError(f"no {family.upper()} controls in {profile.catalog_version_id} "
                         "-- load the 800-53 catalog first")

    if template_path:
        # Clone the user's document: inherit its fonts, heading styles, headers/
        # footers, theme, and page setup; regenerate the body content into it.
        doc = docx.Document(str(template_path))
        _clear_body(doc)
    else:
        doc = docx.Document()
        style = doc.styles["Normal"]
        style.font.name = "Arial"; style.font.size = Pt(11)
    TS = _table_style(doc) or "Table Grid"
    _STYLES = {s.name for s in doc.styles}

    def P(text, style=None):
        """add_paragraph that falls back to default when a style is absent in the template."""
        return doc.add_paragraph(text, style=style if style in _STYLES else None)

    # ---- Title page ----
    title_txt = profile.display_name()
    if profile.system_long_name and profile.system_name:
        title_txt = f"{profile.system_long_name} ({profile.system_name})"
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title_txt)
    run.bold = True; run.font.size = Pt(18); color_run(run, brand_color)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"Security Assessment and Authorization Plan\n"
                     f"{fam_title} ({family.upper()}) Control Family")
    r2.bold = True; r2.font.size = Pt(14)
    prepared_block(doc, prepared_by, profile.agency)
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version 1.0 (DRAFT — needs review)\n{_dt.date.today():%d %b %Y}")

    # Revision history
    doc.add_paragraph()
    P("Revision History", "Heading 3")
    rev = doc.add_table(rows=1, cols=4); rev.style = TS
    for i, h in enumerate(("Date", "Version", "Author", "Changes Made / Section(s)")):
        rev.rows[0].cells[i].text = h
    rc = rev.add_row().cells
    rc[0].text = f"{_dt.date.today():%d %b %Y}"; rc[1].text = "1.0"
    rc[2].text = "ComplyForge (draft)"; rc[3].text = "Initial generated draft"
    doc.add_page_break()

    # ---- 1. Introduction ----
    P("Introduction", "Heading 1")
    doc.add_paragraph(
        f"This control family plan addresses the {fam_title} ({family.upper()}) controls "
        f"for {profile.description()} It documents how the system implements the selected "
        f"security controls and control enhancements.")

    P("Purpose", "Heading 2")
    doc.add_paragraph(
        f"The purpose of this plan is to clearly address the {family.upper()} security "
        f"controls listed in the compliance matrix below, as selected for the "
        f"{profile.baseline} baseline.")
    # Compliance matrix
    cap = doc.add_paragraph(); cap.add_run(
        f"Table 1.1 — NIST SP 800-53 {family.upper()} Compliance Matrix").italic = True
    matrix = doc.add_table(rows=1, cols=3); matrix.style = TS
    for i, h in enumerate(("No.", "Control", profile.baseline)):
        matrix.rows[0].cells[i].text = h
    for n, c in enumerate(controls, 1):
        cells = matrix.add_row().cells
        cells[0].text = str(n)
        cells[1].text = f"{c['control_id'].upper()} — {c['title']}"
        cells[2].text = "X"

    P("Scope", "Heading 2")
    enc = f" operating within the {profile.enclave}" if profile.enclave else ""
    doc.add_paragraph(
        f"The scope of this plan is limited to {profile.short_name()}{enc}, to include "
        f"external system connections and third-party service providers.")

    P("Roles and Responsibilities", "Heading 2")
    cap = doc.add_paragraph(); cap.add_run(
        f"Table 1.2 — {profile.short_name()} Roles and Responsibilities").italic = True
    rt = doc.add_table(rows=1, cols=3); rt.style = TS
    for i, h in enumerate(("Role", "Type", "Responsibilities")):
        rt.rows[0].cells[i].text = h
    for role, typ, resp in profile.roles:
        cells = rt.add_row().cells
        cells[0].text = role; cells[1].text = typ; cells[2].text = resp

    P("Government Personnel", "Heading 2")
    cap = doc.add_paragraph(); cap.add_run("Table 1.3 — Personnel").italic = True
    pt = doc.add_table(rows=1, cols=3); pt.style = TS
    for i, h in enumerate(("Role", "Name", "Organization")):
        pt.rows[0].cells[i].text = h
    for p in profile.personnel:
        cells = pt.add_row().cells
        cells[0].text = p.role; cells[1].text = p.name; cells[2].text = p.org

    P("Applicable Guidance and Directives", "Heading 2")
    for g in profile.guidance:
        P(g, "List Bullet")

    P("Dissemination", "Heading 2")
    doc.add_paragraph(
        f"This document must be made readily available to all personnel supporting "
        f"{profile.short_name()} in a management or privileged function via eMASS.")

    P("Review and Authorization", "Heading 2")
    doc.add_paragraph(
        "This artifact is scheduled to be reviewed on an annual basis in accordance with "
        "the continuous monitoring plan, or whenever a significant change occurs.")
    doc.add_page_break()

    # ---- 2. Family body ----
    P(fam_title.upper(), "Heading 1")
    structured_count = 0
    for c in controls:
        P(c["title"].upper(), "Heading 2")
        narrative, structured = _draft_narrative(provider, c, profile)
        structured_count += int(structured)
        for para in narrative.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        # CCI assessment-procedure table -- populated from loaded DISA CCIs
        cap = doc.add_paragraph()
        cap.add_run(f"{c['control_id'].upper()} Assessment Procedures (CCIs)").italic = True
        try:
            from . import cci as _cci
            rows = _cci.controls_ccis(conn, c["control_id"])
        except Exception:
            rows = []
        tbl = doc.add_table(rows=1, cols=4); tbl.style = TS
        for i, h in enumerate(("AP Acronym", "CCI", "CCI Definition", "Response")):
            tbl.rows[0].cells[i].text = h
        if rows:
            for r in rows:
                cells = tbl.add_row().cells
                cells[0].text = r["ap_acronym"]; cells[1].text = r["cci"]
                cells[2].text = r["definition"]; cells[3].text = ""
        else:
            cells = tbl.add_row().cells
            cells[2].text = "[Load DISA CCIs: python3 scripts/fetch_cci.py]"

    out_path = Path(out_path) if out_path else (
        Path(__file__).resolve().parent.parent / "out" / "control_plans" /
        f"{profile.system_name}_{family.upper()}_Plan.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
